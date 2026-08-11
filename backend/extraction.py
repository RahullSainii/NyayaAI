"""Extract plain text from uploaded documents (PDF, DOCX, and text formats).

Images are not supported because OCR requires a Tesseract binary that isn't
available in the deployment environment. Callers should surface the ValueError
messages to the user.
"""

import io
import logging
from typing import Dict

logger = logging.getLogger(__name__)

# Cap the returned text so a huge document doesn't blow up the prompt / payload.
MAX_EXTRACT_CHARS = 20000

_TEXT_EXTS = (
    ".txt", ".md", ".markdown", ".csv", ".json", ".log", ".rtf",
    ".html", ".htm", ".xml", ".yaml", ".yml",
)
_IMAGE_EXTS = (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tif", ".tiff")


def _extract_pdf(data: bytes) -> str:
    import pdfplumber

    parts = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
            try:
                page.flush_cache()
            except Exception:
                pass
            # Stop early once we clearly have more than we'll keep.
            if sum(len(p) for p in parts) > MAX_EXTRACT_CHARS * 2:
                break
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    blocks = [p.text for p in document.paragraphs]
    # Include table cell text, common in legal documents.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n".join(b for b in blocks if b is not None)


def _extract_plain(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="replace")


def extract_text(filename: str, data: bytes) -> Dict:
    """Return {"text", "truncated"} or raise ValueError with a user-facing reason."""
    name = (filename or "").lower().strip()

    if name.endswith(".pdf"):
        text = _extract_pdf(data)
    elif name.endswith(".docx"):
        text = _extract_docx(data)
    elif name.endswith(_TEXT_EXTS):
        text = _extract_plain(data)
    elif name.endswith(_IMAGE_EXTS):
        raise ValueError(
            "Images aren't supported for text extraction. Please attach a PDF, DOCX, or text document."
        )
    elif name.endswith(".doc"):
        raise ValueError("Legacy .doc files aren't supported. Please save it as .docx or PDF.")
    else:
        raise ValueError("Unsupported file type. Attach a PDF, DOCX, or text document.")

    text = (text or "").strip()
    if not text:
        raise ValueError(
            "No readable text found — the file may be scanned or image-based (which needs OCR)."
        )

    return {"text": text[:MAX_EXTRACT_CHARS], "truncated": len(text) > MAX_EXTRACT_CHARS}
